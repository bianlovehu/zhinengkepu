"""
文档加载器
"""
import json
import logging
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """文档数据结构"""
    content: str
    source: str
    metadata: Dict[str, Any]


class DocumentLoader:
    """
    文档加载器

    支持多种格式：TXT, PDF, DOCX, HTML, Markdown
    """

    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".html", ".md", ".markdown"}
    MANUAL_BOUNDARY = "\n\n<<<MANUAL_BOUNDARY>>>\n\n"

    def __init__(self):
        self.parsers = {
            ".txt": self._parse_txt,
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".html": self._parse_html,
            ".md": self._parse_md,
        }

    def load_file(self, file_path: str) -> Document:
        """
        加载单个文件

        Args:
            file_path: 文件路径

        Returns:
            Document: 文档对象
        """
        documents = self.load_documents_from_file(file_path)
        if not documents:
            raise ValueError(f"No document content loaded: {file_path}")
        return documents[0]

    def load_documents_from_file(self, file_path: str) -> List[Document]:
        """加载单个文件，必要时拆成多个逻辑文档。"""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        # 获取解析器
        parser = self.parsers.get(ext, self._parse_txt)

        # 从文件名提取文档类型/标题作为上下文
        # 例如："电钻手册.txt" -> "电钻"
        # 例如："发电机使用指南.pdf" -> "发电机"
        title = self._extract_title_from_filename(path.name)

        content = parser(path)
        parts = self._split_logical_documents(content)
        documents = []

        for idx, part in enumerate(parts):
            if not part.strip():
                continue

            logical_title = self._infer_logical_title(part, title, idx)
            manual_id = f"{path.stem}_{idx:03d}" if len(parts) > 1 else path.stem
            source = path.name if len(parts) == 1 else f"{path.stem}#{idx + 1}{path.suffix}"
            documents.append(Document(
                content=part,
                source=source,
                metadata={
                    "path": str(path.absolute()),
                    "type": ext[1:],
                    "size": path.stat().st_size,
                    "title": logical_title,
                    "full_title": logical_title if len(parts) > 1 else path.stem,
                    "manual_id": manual_id,
                    "logical_index": idx,
                    "logical_count": len(parts),
                    "language": self._detect_language(part),
                    "original_source": path.name,
                }
            ))

        return documents

    def load_directory(self, dir_path: str, recursive: bool = True) -> List[Document]:
        """
        加载目录下所有支持的文档

        Args:
            dir_path: 目录路径
            recursive: 是否递归子目录

        Returns:
            List[Document]: 文档列表
        """
        path = Path(dir_path)

        if not path.exists() or not path.is_dir():
            raise NotADirectoryError(f"Invalid directory: {dir_path}")

        documents = []
        pattern = "**/*" if recursive else "*"

        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    docs = self.load_documents_from_file(str(file_path))
                    documents.extend(docs)
                    logger.info(f"Loaded: {file_path.name} ({len(docs)} logical document(s))")
                except Exception as e:
                    logger.error(f"Failed to load {file_path}: {e}")

        logger.info(f"Loaded {len(documents)} documents from {dir_path}")
        return documents

    def _parse_txt(self, path: Path) -> str:
        """解析TXT文件"""
        with open(path, "r", encoding="utf-8") as f:
            raw = f.read()
        return self._normalize_text(self._decode_structured_text(raw))

    def _parse_pdf(self, path: Path) -> str:
        """解析PDF文件"""
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                import PyPDF2 as PdfReader
            except ImportError:
                logger.error("Please install pypdf or PyPDF2")
                return ""

        text = []
        reader = PdfReader(path)
        for page in reader.pages:
            text.append(page.extract_text())

        return "\n".join(text)

    def _parse_docx(self, path: Path) -> str:
        """解析DOCX文件"""
        try:
            from docx import Document
        except ImportError:
            logger.error("Please install python-docx")
            return ""

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs)

    def _parse_html(self, path: Path) -> str:
        """解析HTML文件"""
        from bs4 import BeautifulSoup

        with open(path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f.read(), "html.parser")

        # 移除脚本和样式
        for tag in soup(["script", "style"]):
            tag.decompose()

        return soup.get_text(separator="\n", strip=True)

    def _parse_md(self, path: Path) -> str:
        """解析Markdown文件"""
        with open(path, "r", encoding="utf-8") as f:
            raw = f.read()
        return self._normalize_text(self._decode_structured_text(raw))

    def _decode_structured_text(self, raw: str) -> str:
        """解析 JSON 字符串/列表或逐行 JSON 列表，避免把转义换行当普通字符。"""
        stripped = raw.strip()
        if not stripped:
            return ""

        try:
            return self._coerce_json_value(json.loads(stripped), joiner="\n\n")
        except json.JSONDecodeError:
            pass

        stream_segments = self._decode_json_stream(stripped)
        if len(stream_segments) > 1:
            return self.MANUAL_BOUNDARY.join(stream_segments)

        json_segments = []
        parseable_lines = 0
        for line in stripped.splitlines():
            candidate = line.strip()
            if not candidate:
                continue
            try:
                json_segments.append(self._coerce_json_value(json.loads(candidate), joiner="\n\n"))
                parseable_lines += 1
            except json.JSONDecodeError:
                json_segments = []
                break

        if json_segments and parseable_lines > 1:
            return self.MANUAL_BOUNDARY.join(json_segments)

        # 兼容已经被转成普通文本但仍保留反斜杠转义的文件。
        return stripped

    def _decode_json_stream(self, text: str) -> List[str]:
        """解析连续 JSON 值，例如一行一个数组或多个数组直接拼接。"""
        decoder = json.JSONDecoder()
        idx = 0
        segments = []
        length = len(text)

        while idx < length:
            while idx < length and text[idx].isspace():
                idx += 1
            if idx >= length:
                break
            try:
                value, next_idx = decoder.raw_decode(text, idx)
            except json.JSONDecodeError:
                return segments
            segments.append(self._coerce_json_value(value, joiner="\n\n"))
            idx = next_idx

        return segments

    def _coerce_json_value(self, value: Any, joiner: str = "\n") -> str:
        if isinstance(value, list):
            return joiner.join(self._coerce_json_value(item, joiner="\n") for item in value)
        if isinstance(value, dict):
            return json.dumps(value, ensure_ascii=False)
        return str(value)

    def _normalize_text(self, text: str) -> str:
        """做轻量清洗，保留手册语义结构。"""
        if not text:
            return ""

        if self.MANUAL_BOUNDARY in text:
            return self.MANUAL_BOUNDARY.join(
                self._normalize_text(part)
                for part in text.split(self.MANUAL_BOUNDARY)
                if part.strip()
            )

        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = text.replace("\\n", "\n").replace("\\t", " ")
        text = text.replace('\\"', '"').replace("\\/", "/")
        text = text.replace("\u3000", " ")

        # 把内联标题拉回行首，帮助后续按标题切块。
        text = re.sub(r"(?<!\n)\s+(#{1,6}\s+)", r"\n\1", text)
        text = re.sub(r"\s*(<PIC>(?:[A-Za-z0-9_\-]+)?)\s*", r" \1 ", text)
        text = re.sub(r"[.·。]{6,}", " ", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()

    def _split_logical_documents(self, content: str) -> List[str]:
        parts = [part.strip() for part in content.split(self.MANUAL_BOUNDARY)]
        return [part for part in parts if part]

    def _infer_logical_title(self, content: str, fallback: str, index: int) -> str:
        for line in content.splitlines()[:30]:
            stripped = line.strip(" #\t")
            if not stripped or stripped.startswith("<PIC>"):
                continue
            if 4 <= len(stripped) <= 80:
                return stripped
        return f"{fallback}-{index + 1}" if index else fallback

    def _detect_language(self, content: str) -> str:
        sample = content[:2000]
        chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", sample))
        ascii_letters = len(re.findall(r"[A-Za-z]", sample))
        if chinese_chars and chinese_chars >= ascii_letters * 0.15:
            return "zh"
        if ascii_letters:
            return "en"
        return "unknown"

    def _extract_title_from_filename(self, filename: str) -> str:
        """
        从文件名提取文档标题/类型

        例如：
        - "电钻手册.txt" -> "电钻"
        - "发电机使用指南.pdf" -> "发电机"
        - "DCB107充电说明.docx" -> "DCB107充电"
        """
        import re

        # 移除扩展名
        name = Path(filename).stem

        # 常见的产品类型词
        product_keywords = [
            "电钻", "充电器", "电池", "电池包", "电动工具",
            "发电机", "洗碗机", "VR", "头显", "显示器",
            "电锯", "角磨机", "冲击钻", "螺丝刀", "锯子",
            "手册", "说明书", "指南", "使用说明", "操作指南"
        ]

        # 尝试提取产品类型
        for keyword in product_keywords:
            if keyword in name:
                # 如果包含产品词，截取产品词及其后的描述
                idx = name.find(keyword)
                title = name[idx:]
                # 移除"手册"、"说明书"等后缀，保留核心产品名
                for suffix in ["手册", "说明书", "指南", "使用说明", "操作指南"]:
                    title = title.replace(suffix, "").strip()
                if title:
                    return title
                return keyword

        # 如果没有匹配，返回清理后的文件名
        return re.sub(r'[_\-]?(手册|说明书|指南|使用说明|操作指南)$', '', name).strip() or name

    def extract_images_from_pdf(self, pdf_path: str, output_dir: str) -> List[Dict[str, Any]]:
        """
        从PDF中提取图片

        Args:
            pdf_path: PDF文件路径
            output_dir: 图片输出目录

        Returns:
            List[Dict]: 图片信息列表
        """
        images = []
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("Please install pymupdf for PDF image extraction")
            return images

        doc = fitz.open(pdf_path)
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        for page_num, page in enumerate(doc):
            image_list = page.get_images(full=True)
            for img_idx, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]

                image_name = f"{Path(pdf_path).stem}_p{page_num+1}_img{img_idx+1}.{image_ext}"
                image_path = output_path / image_name

                with open(image_path, "wb") as f:
                    f.write(image_bytes)

                images.append({
                    "id": f"{Path(pdf_path).stem}_p{page_num+1}_{img_idx+1}",
                    "path": str(image_path),
                    "page": page_num + 1,
                    "ext": image_ext
                })

        doc.close()
        logger.info(f"Extracted {len(images)} images from {pdf_path}")
        return images
